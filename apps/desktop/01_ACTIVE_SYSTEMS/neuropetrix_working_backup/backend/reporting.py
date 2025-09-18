from typing import List, Dict
from rapidfuzz import fuzz, process
import io, re
from docx import Document
from docx.shared import Pt

def sentence_diff(prev_text: str, curr_text: str) -> List[Dict]:
    """İki metin arasında cümle bazında fark analizi"""
    prev = [s.strip() for s in re.split(r'(?<=[\.\?\!])\s+', (prev_text or "").strip()) if s.strip()]
    curr = [s.strip() for s in re.split(r'(?<=[\.\?\!])\s+', (curr_text or "").strip()) if s.strip()]
    out = []
    
    # Mevcut cümleleri analiz et
    for c in curr:
        m, score, _ = process.extractOne(c, prev, scorer=fuzz.token_set_ratio) if prev else (None, 0, None)
        if score >= 85:
            out.append({
                "type": "same" if c == m else "changed",
                "old": m or "",
                "new": c,
                "score": float(score)
            })
        else:
            out.append({
                "type": "added",
                "old": "",
                "new": c,
                "score": float(score)
            })
    
    # Önceki cümlelerde kaldırılanları bul
    for p in prev:
        m, score, _ = process.extractOne(p, curr, scorer=fuzz.token_set_ratio) if curr else (None, 0, None)
        if score < 85:
            out.append({
                "type": "removed",
                "old": p,
                "new": "",
                "score": float(score)
            })
    
    return out

def build_docx(title: str, sections: Dict[str, str]) -> bytes:
    """Markdown metni DOCX formatına dönüştür"""
    d = Document()
    s = d.styles['Normal']
    s.font.name = 'Calibri'
    s.font.size = Pt(11)
    
    d.add_heading(title or 'NeuroPETrix Report', level=1)
    
    for sec, txt in sections.items():
        if txt and txt.strip():
            d.add_heading(sec, level=2)
            for line in txt.split('\n'):
                d.add_paragraph(line.strip())
    
    bio = io.BytesIO()
    d.save(bio)
    return bio.getvalue()

def ebm_composite_score(grade_level: str, rob2: Dict[str, int]) -> int:
    """GRADE + RoB2 birleşik skor (0-100 basit kural tabanlı)"""
    # GRADE seviyesine göre temel skor
    base = {
        "high": 90,
        "moderate": 70,
        "low": 50,
        "very low": 30
    }.get(grade_level.lower(), 50)
    
    # RoB2 cezaları (0 iyi, 1 orta risk, 2 yüksek risk)
    penalty = 0
    for k, v in rob2.items():
        penalty += v * 10
    
    # Final skor (0-100 arası)
    return max(0, min(100, base - penalty))

def format_diff_html(diff_results: List[Dict]) -> str:
    """Diff sonuçlarını HTML formatında döndür"""
    html = "<div class='diff-container'>"
    
    for item in diff_results:
        if item["type"] == "same":
            html += f"<div class='diff-same'>{item['new']}</div>"
        elif item["type"] == "changed":
            html += f"<div class='diff-changed'>"
            html += f"<span class='diff-old'>{item['old']}</span>"
            html += f" → <span class='diff-new'>{item['new']}</span>"
            html += f" <span class='diff-score'>({item['score']:.1f}% benzer)</span>"
            html += "</div>"
        elif item["type"] == "added":
            html += f"<div class='diff-added'>+ {item['new']}</div>"
        elif item["type"] == "removed":
            html += f"<div class='diff-removed'>- {item['old']}</div>"
    
    html += "</div>"
    return html

def extract_tsnm_metrics(text: str) -> Dict[str, float]:
    """TSNM metninden sayısal metrikleri çıkar"""
    metrics = {}
    
    # SUVmax pattern
    suv_match = re.search(r'SUVmax[:\s]*([0-9]+\.?[0-9]*)', text, re.IGNORECASE)
    if suv_match:
        metrics['suvmax'] = float(suv_match.group(1))
    
    # MTV pattern
    mtv_match = re.search(r'MTV[:\s]*([0-9]+\.?[0-9]*)', text, re.IGNORECASE)
    if mtv_match:
        metrics['mtv'] = float(mtv_match.group(1))
    
    # TLG pattern
    tlg_match = re.search(r'TLG[:\s]*([0-9]+\.?[0-9]*)', text, re.IGNORECASE)
    if tlg_match:
        metrics['tlg'] = float(tlg_match.group(1))
    
    return metrics

def calculate_progression_flag(prev_metrics: Dict[str, float], curr_metrics: Dict[str, float]) -> Dict:
    """SUVmax değişimine göre progresyon flag'i hesapla"""
    if 'suvmax' not in prev_metrics or 'suvmax' not in curr_metrics:
        return {"type": "unknown", "change_percent": None}
    
    prev_suv = prev_metrics['suvmax']
    curr_suv = curr_metrics['suvmax']
    
    if prev_suv <= 0:
        return {"type": "unknown", "change_percent": None}
    
    change_percent = ((curr_suv - prev_suv) / prev_suv) * 100
    
    if change_percent >= 30:
        return {"type": "progression", "change_percent": round(change_percent, 1)}
    elif change_percent <= -30:
        return {"type": "response", "change_percent": round(change_percent, 1)}
    else:
        return {"type": "stable", "change_percent": round(change_percent, 1)}
