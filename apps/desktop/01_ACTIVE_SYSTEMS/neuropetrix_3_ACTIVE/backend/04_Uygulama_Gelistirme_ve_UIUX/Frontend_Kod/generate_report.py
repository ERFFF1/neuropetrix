from typing import Dict, Any
from datetime import datetime
import json

def generate_report(report_type: str, data: Dict[str, Any]) -> str:
    """Rapor oluştur"""
    
    if report_type == "tsnm":
        return generate_tsnm_report(data)
    elif report_type == "standard":
        return generate_standard_report(data)
    elif report_type == "ai_enhanced":
        return generate_ai_enhanced_report(data)
    else:
        return generate_standard_report(data)

def generate_tsnm_report(data: Dict[str, Any]) -> str:
    """TSNM formatında rapor oluştur"""
    
    template = f"""# {data.get('modality', 'PET/CT')} Raporu

## Hasta Bilgileri
- **Ad Soyad:** {data.get('patient', {}).get('name', 'N/A')}
- **Yaş:** {data.get('patient', {}).get('age', 'N/A')}
- **Cinsiyet:** {data.get('patient', {}).get('gender', 'N/A')}
- **Hasta No:** {data.get('patient', {}).get('id', 'N/A')}

## İndikasyon
{data.get('indication', 'Belirtilmemiş')}

## Teknik Bilgiler
- **Cihaz:** {data.get('device_model', 'Siemens Biograph mCT')}
- **Doza:** {data.get('dose_mbq', '185')} MBq
- **Glikoz:** {data.get('glycemia_mgdl', '110')} mg/dL
- **Açlık Süresi:** {data.get('fasting_hours', '6')} saat

## Bulgular
{data.get('findings', 'Bulgular henüz girilmemiş.')}

## SUV Değerleri
- **Lezyon SUVmax:** {data.get('suv_max', 'N/A')}
- **Lezyon SUVmean:** {data.get('suv_mean', 'N/A')}
- **Karaciğer SUVmean:** {data.get('liver_suv', '2.1')} (Referans)

## Sonuç
{data.get('conclusion', 'Sonuç henüz girilmemiş.')}

## Öneriler
{data.get('recommendations', 'Öneriler henüz girilmemiş.')}

---
*Rapor {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} tarihinde oluşturuldu.*
*TSNM Kılavuzlarına Uygun - NeuroPETrix AI Sistemi*
"""
    
    return template

def generate_standard_report(data: Dict[str, Any]) -> str:
    """Standart rapor oluştur"""
    
    template = f"""# Klinik Rapor

## Hasta Bilgileri
- **Ad:** {data.get('patient', {}).get('name', 'N/A')}
- **Yaş:** {data.get('patient', {}).get('age', 'N/A')}
- **Cinsiyet:** {data.get('patient', {}).get('gender', 'N/A')}

## Bulgular
{data.get('findings', 'Bulgular henüz girilmemiş.')}

## Sonuç
{data.get('conclusion', 'Sonuç henüz girilmemiş.')}

---
*Rapor {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} tarihinde oluşturuldu.*
"""
    
    return template

def generate_ai_enhanced_report(data: Dict[str, Any]) -> str:
    """AI destekli rapor oluştur"""
    
    template = f"""# AI Destekli PET/CT Raporu

## Hasta Bilgileri
- **Ad Soyad:** {data.get('patient', {}).get('name', 'N/A')}
- **Yaş:** {data.get('patient', {}).get('age', 'N/A')}
- **Cinsiyet:** {data.get('patient', {}).get('gender', 'N/A')}

## AI Analiz Sonuçları
- **Güven Skoru:** {data.get('ai_confidence', 'N/A')}
- **Tespit Edilen Lezyon:** {data.get('lesions_detected', 'N/A')}
- **Segmentasyon Kalitesi:** {data.get('segmentation_quality', 'N/A')}

## Bulgular
{data.get('findings', 'Bulgular henüz girilmemiş.')}

## AI Önerileri
{data.get('ai_recommendations', 'AI önerisi bulunamadı.')}

## Sonuç
{data.get('conclusion', 'Sonuç henüz girilmemiş.')}

---
*Rapor {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} tarihinde oluşturuldu.*
*NeuroPETrix AI Sistemi ile desteklenmiştir.*
"""
    
    return template

def export_report_to_docx(report_content: str, filename: str = None) -> bytes:
    """Raporu DOCX formatında dışa aktar"""
    try:
        from docx import Document
        from docx.shared import Inches
        
        doc = Document()
        doc.add_heading('NeuroPETrix Raporu', 0)
        
        # İçeriği paragraflar halinde ekle
        for line in report_content.split('\n'):
            if line.strip():
                if line.startswith('#'):
                    # Başlık
                    level = line.count('#')
                    text = line.lstrip('#').strip()
                    doc.add_heading(text, level)
                else:
                    # Normal paragraf
                    doc.add_paragraph(line)
        
        # Dosyayı kaydet
        if filename:
            doc.save(filename)
            return None
        else:
            # Bytes olarak döndür
            import io
            buffer = io.BytesIO()
            doc.save(buffer)
            return buffer.getvalue()
            
    except ImportError:
        raise Exception("python-docx kütüphanesi gerekli")

def export_report_to_pdf(report_content: str, filename: str = None) -> bytes:
    """Raporu PDF formatında dışa aktar"""
    try:
        import markdown
        from weasyprint import HTML
        
        # Markdown'ı HTML'e çevir
        html_content = markdown.markdown(report_content)
        
        # HTML'i PDF'e çevir
        html = HTML(string=html_content)
        pdf = html.write_pdf()
        
        if filename:
            with open(filename, 'wb') as f:
                f.write(pdf)
            return None
        else:
            return pdf
            
    except ImportError:
        raise Exception("markdown ve weasyprint kütüphaneleri gerekli")


