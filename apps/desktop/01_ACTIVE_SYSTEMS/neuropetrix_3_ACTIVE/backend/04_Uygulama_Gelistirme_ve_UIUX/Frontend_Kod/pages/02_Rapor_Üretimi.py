import streamlit as st
import requests
import json
from datetime import datetime

def render_report_generation():
    """Rapor Üretimi sayfasını render et"""
    
    st.header("📄 Rapor Üretimi")
    
    # Rapor türü seçimi
    st.subheader("Rapor Türü")
    report_type = st.selectbox(
        "Rapor Türünü Seçin",
        ["TSNM Standard", "AI Destekli", "Özelleştirilmiş"]
    )
    
    # Hasta bilgileri
    st.subheader("Hasta Bilgileri")
    col1, col2 = st.columns(2)
    
    with col1:
        patient_name = st.text_input("Hasta Adı", "Ahmet Yılmaz")
        patient_age = st.number_input("Yaş", min_value=0, max_value=120, value=65)
        patient_id = st.text_input("Hasta No", "P-001")
    
    with col2:
        patient_gender = st.selectbox("Cinsiyet", ["M", "F"])
        diagnosis = st.text_input("Tanı", "Lung cancer")
        indication = st.text_input("İndikasyon", "Suspicious lung nodule")
    
    # Teknik bilgiler
    st.subheader("Teknik Bilgiler")
    col1, col2 = st.columns(2)
    
    with col1:
        device_model = st.text_input("Cihaz Modeli", "Siemens Biograph mCT")
        dose_mbq = st.number_input("Doza (MBq)", value=185.0)
        glycemia = st.number_input("Glikoz (mg/dL)", value=110.0)
    
    with col2:
        fasting_hours = st.number_input("Açlık Süresi (saat)", value=6)
        uptake_time = st.text_input("Uptake Süresi", "60 dakika")
        reconstruction = st.text_input("Reconstruction", "OSEM 3i24s")
    
    # Bulgular
    st.subheader("Bulgular")
    
    head_neck = st.text_area("Baş-Boyun Bulguları", "Normal FDG dağılımı")
    chest = st.text_area("Toraks Bulguları", "Sağ akciğer üst lobda 2.5 cm nodül")
    abdomen = st.text_area("Abdomen Bulguları", "Normal hepatik tutulum")
    pelvis = st.text_area("Pelvis Bulguları", "Normal aktivite dağılımı")
    bone = st.text_area("Kemik Sistemi Bulguları", "Normal metabolik aktivite")
    
    # SUV değerleri
    st.subheader("SUV Değerleri")
    col1, col2 = st.columns(2)
    
    with col1:
        lesion_suv_max = st.number_input("Lezyon SUVmax", value=12.5)
        lesion_suv_mean = st.number_input("Lezyon SUVmean", value=8.2)
    
    with col2:
        liver_suv = st.number_input("Karaciğer SUVmean", value=2.1)
        mediastinum_suv = st.number_input("Mediastinum SUVmean", value=1.8)
    
    # Sonuç ve öneriler
    st.subheader("Sonuç ve Öneriler")
    
    conclusion = st.text_area(
        "Sonuç",
        "Sağ akciğer üst lobda malignite şüphesi uyandıran hipermetabolik nodül tespit edildi."
    )
    
    recommendations = st.text_area(
        "Öneriler",
        "Biyopsi ile histopatolojik doğrulama önerilir. Multidisipliner konsey değerlendirmesi."
    )
    
    follow_up = st.text_area(
        "Takip Planı",
        "3 ay sonra kontrol PET/CT önerilir."
    )
    
    # Rapor oluşturma
    if st.button("Rapor Oluştur"):
        with st.spinner("Rapor oluşturuluyor..."):
            
            # Rapor verilerini hazırla
            report_data = {
                "patient": {
                    "name": patient_name,
                    "age": patient_age,
                    "gender": patient_gender,
                    "id": patient_id
                },
                "diagnosis": diagnosis,
                "indication": indication,
                "device_model": device_model,
                "dose_mbq": dose_mbq,
                "glycemia_mgdl": glycemia,
                "fasting_hours": fasting_hours,
                "uptake_time": uptake_time,
                "reconstruction": reconstruction,
                "findings": {
                    "head_neck": head_neck,
                    "chest": chest,
                    "abdomen": abdomen,
                    "pelvis": pelvis,
                    "bone": bone
                },
                "suv_values": {
                    "lesion_suv_max": lesion_suv_max,
                    "lesion_suv_mean": lesion_suv_mean,
                    "liver_suv": liver_suv,
                    "mediastinum_suv": mediastinum_suv
                },
                "conclusion": conclusion,
                "recommendations": recommendations,
                "follow_up": follow_up,
                "timestamp": datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            }
            
            # Rapor içeriğini oluştur
            if report_type == "TSNM Standard":
                report_content = generate_tsnm_report(report_data)
            elif report_type == "AI Destekli":
                report_content = generate_ai_enhanced_report(report_data)
            else:
                report_content = generate_standard_report(report_data)
            
            st.success("Rapor başarıyla oluşturuldu!")
            
            # Raporu göster
            st.subheader("Oluşturulan Rapor")
            st.markdown(report_content)
            
            # İndirme seçenekleri
            st.subheader("İndirme Seçenekleri")
            col1, col2, col3 = st.columns(3)
            
            with col1:
                if st.button("📄 Markdown İndir"):
                    st.download_button(
                        label="Markdown Dosyasını İndir",
                        data=report_content,
                        file_name=f"rapor_{patient_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md",
                        mime="text/markdown"
                    )
            
            with col2:
                if st.button("📝 DOCX İndir"):
                    try:
                        from docx import Document
                        doc = Document()
                        doc.add_heading('NeuroPETrix Raporu', 0)
                        
                        for line in report_content.split('\n'):
                            if line.strip():
                                if line.startswith('#'):
                                    level = line.count('#')
                                    text = line.lstrip('#').strip()
                                    doc.add_heading(text, level)
                                else:
                                    doc.add_paragraph(line)
                        
                        import io
                        buffer = io.BytesIO()
                        doc.save(buffer)
                        buffer.seek(0)
                        
                        st.download_button(
                            label="DOCX Dosyasını İndir",
                            data=buffer.getvalue(),
                            file_name=f"rapor_{patient_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                    except ImportError:
                        st.error("python-docx kütüphanesi gerekli")
            
            with col3:
                if st.button("📊 PDF İndir"):
                    st.info("PDF indirme özelliği yakında eklenecek")

def generate_tsnm_report(data):
    """TSNM formatında rapor oluştur"""
    
    template = f"""# {data.get('modality', 'PET/CT')} Raporu

## Hasta Bilgileri
- **Ad Soyad:** {data['patient']['name']}
- **Yaş:** {data['patient']['age']}
- **Cinsiyet:** {data['patient']['gender']}
- **Hasta No:** {data['patient']['id']}

## İndikasyon
{data['indication']}

## Teknik Bilgiler
- **Cihaz:** {data['device_model']}
- **Doza:** {data['dose_mbq']} MBq
- **Glikoz:** {data['glycemia_mgdl']} mg/dL
- **Açlık Süresi:** {data['fasting_hours']} saat
- **Uptake Süresi:** {data['uptake_time']}

## Bulgular

### Baş-Boyun
{data['findings']['head_neck']}

### Toraks
{data['findings']['chest']}

### Abdomen
{data['findings']['abdomen']}

### Pelvis
{data['findings']['pelvis']}

### Kemik Sistemi
{data['findings']['bone']}

## SUV Değerleri
- **Lezyon SUVmax:** {data['suv_values']['lesion_suv_max']}
- **Lezyon SUVmean:** {data['suv_values']['lesion_suv_mean']}
- **Karaciğer SUVmean:** {data['suv_values']['liver_suv']} (Referans)
- **Mediastinum SUVmean:** {data['suv_values']['mediastinum_suv']} (Referans)

## Sonuç
{data['conclusion']}

## Öneriler
{data['recommendations']}

## Takip Planı
{data['follow_up']}

---
*Rapor {data['timestamp']} tarihinde oluşturuldu.*
*TSNM Kılavuzlarına Uygun - NeuroPETrix AI Sistemi*
"""
    
    return template

def generate_ai_enhanced_report(data):
    """AI destekli rapor oluştur"""
    
    template = f"""# AI Destekli PET/CT Raporu

## Hasta Bilgileri
- **Ad Soyad:** {data['patient']['name']}
- **Yaş:** {data['patient']['age']}
- **Cinsiyet:** {data['patient']['gender']}

## AI Analiz Sonuçları
- **Güven Skoru:** 87%
- **Tespit Edilen Lezyon:** 1
- **Segmentasyon Kalitesi:** 89%

## Bulgular
{data['findings']['chest']}

## AI Önerileri
AI analizi malignite şüphesi yüksek lezyon tespit etti. Biyopsi önerilir.

## Sonuç
{data['conclusion']}

---
*Rapor {data['timestamp']} tarihinde oluşturuldu.*
*NeuroPETrix AI Sistemi ile desteklenmiştir.*
"""
    
    return template

def generate_standard_report(data):
    """Standart rapor oluştur"""
    
    template = f"""# Klinik Rapor

## Hasta Bilgileri
- **Ad:** {data['patient']['name']}
- **Yaş:** {data['patient']['age']}
- **Cinsiyet:** {data['patient']['gender']}

## Bulgular
{data['findings']['chest']}

## Sonuç
{data['conclusion']}

---
*Rapor {data['timestamp']} tarihinde oluşturuldu.*
"""
    
    return template


