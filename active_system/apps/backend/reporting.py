from typing import List, Dict, Any
from datetime import datetime
import json

def sentence_diff(old_sentences: List[str], new_sentences: List[str]) -> Dict[str, Any]:
    """Cümle farklarını hesapla"""
    added = [s for s in new_sentences if s not in old_sentences]
    removed = [s for s in old_sentences if s not in new_sentences]
    
    return {
        "added": added,
        "removed": removed,
        "unchanged": [s for s in old_sentences if s in new_sentences]
    }

def build_docx(content: str, filename: str = None) -> bytes:
    """DOCX dosyası oluştur"""
    from docx import Document
    from docx.shared import Inches
    
    doc = Document()
    doc.add_heading('NeuroPETrix Raporu', 0)
    
    # İçeriği paragraflar halinde ekle
    for line in content.split('\n'):
        if line.strip():
            if line.startswith('#'):
                # Başlık
                level = line.count('#')
                text = line.lstrip('#').strip()
                doc.add_heading(text, level)
            else:
                # Normal paragraf
                doc.add_paragraph(line)
    
    # Dosyayı kaydet
    if filename:
        doc.save(filename)
        return None
    else:
        # Bytes olarak döndür
        import io
        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

def ebm_composite_score(evidence_levels: List[str]) -> float:
    """Kanıta dayalı tıp kompozit skoru hesapla"""
    level_scores = {
        '1A': 1.0,
        '1B': 0.9,
        '2A': 0.8,
        '2B': 0.7,
        '3': 0.5,
        'Expert': 0.6,
        'Good Practice': 0.7
    }
    
    if not evidence_levels:
        return 0.0
    
    scores = [level_scores.get(level, 0.5) for level in evidence_levels]
    return sum(scores) / len(scores)

def format_diff_html(diff_result: Dict[str, Any]) -> str:
    """Fark sonuçlarını HTML formatında döndür"""
    html = "<div class='diff-container'>"
    
    if diff_result.get("added"):
        html += "<div class='diff-added'><h4>Eklenen:</h4><ul>"
        for item in diff_result["added"]:
            html += f"<li>{item}</li>"
        html += "</ul></div>"
    
    if diff_result.get("removed"):
        html += "<div class='diff-removed'><h4>Çıkarılan:</h4><ul>"
        for item in diff_result["removed"]:
            html += f"<li>{item}</li>"
        html += "</ul></div>"
    
    html += "</div>"
    return html
